import os
import requests
import datetime
import asyncio
from bs4 import BeautifulSoup
from docx import Document

#配置
class CSaveConfig:
    mode = 0
    def __init__(self,mode = 0):
        CSaveConfig.mode = mode 
    
    @staticmethod
    def GetPathAndName():
        "获取路径和文件名"
        if CSaveConfig.mode == 0:
            return CSaveConfig.__getPathAndName()
    @staticmethod
    def __getDirFileCount(path):
        "获取目录下文件数量"
        file_count = 0
        file_list=os.listdir(path)
        for i in file_list:
            path_now = path + "\\" + i
            if os.path.isdir(path_now) == False:
                file_count = file_count + 1
        return file_count

    @staticmethod  
    def __getPathAndName():
        "路径: CNN_News/data/yyyy/mm/dd/"
        tody = datetime.datetime.now()
        path = '{0}/data/{1}/{2}/{3}/'.format(os.getcwd(), tody.year, tody.month, tody.day)
        if not os.path.exists(path):
            os.makedirs(path)
        
        file_count = CSaveConfig.__getDirFileCount(path)
        name = '{0}-{1}'.format(tody.strftime('%Y%m%d'),file_count+1)
        
        return path,name
    
#保存为文件
class CSaveData:
    def __init__(self,title,content,ext="txt"):
        if ext == "txt":
            self.__toTxt(title,content)
        elif ext == "docx":
            self.__toDocx(title,content)
        else:
            self.__toCustom(content,ext)

    def __toDocx(self,title,content):
        "保存为docx文件"
        # 创建一个文档对象
        doc = Document()

        # 添加标题
        doc.add_heading(title, level=1)

        # 添加段落
        for p in content:
            doc.add_paragraph(p)

        "函数返回参数是元组，这里用*进行解包"
        file = "{0}{1}.docx".format(*CSaveConfig.GetPathAndName())
        # 保存文档
        doc.save(file)

    def __toTxt(self,title,content):
        "保存为txt文件"
        file = "{0}{1}.txt".format(*CSaveConfig.GetPathAndName())
        with open(file, 'w', encoding='utf-8') as f:
            if title is not None:
                f.write(title)
                f.write('\n')
            for line in content:
                f.write(line)
                f.write('\n')
            

    def __toCustom(self,content,ext):
        "保存为自定义格式文件"
        if type(content)!= bytes:
            content = content.encode("utf-8")
        file = "{0}{1}.{2}".format(*CSaveConfig.GetPathAndName(),ext)
        with open(file, 'wb') as f:
            f.write(content)


#获取新闻
class CGetNews:
    def __init__(self,domain=""):
        self.domain = domain
        if self.domain[-1] == "/":
            self.domain = self.domain[:-1]

    def Get(self):
        result = requests.get("{0}/world".format(self.domain))
        if result is None: 
            raise Exception("获取列表失败[{0}]".format(self.domain))
        if result.status_code != 200:
            raise Exception("获取列表失败[{0}][{1}]".format(self.domain,result.status_code))
        "1、获取列表页"
        'class="container__link container__link--type-article container_lead-plus-headlines__link"'
        news_list = set()
        soup = BeautifulSoup(result.content, 'html.parser')
        links = soup.select('a[class="container__link container__link--type-article container_lead-plus-headlines__link"]')
        print("今日文章数:{0}".format(len(links)))
        for link in links:
            href = link.get('href')
            if href is None:
                continue
            "/2024/12/24/sport/caitlin-clark-ap-female-athlete-of-the-year-spt/index.html"
            news_list.add("{0}{1}".format(self.domain,href))

        "2、获取详情页"
        for url in news_list:
            asyncio.run(self.__getArticle(url))
            #break
        #CSaveData("我是标题",content,"html")
        

    async def __getArticle(self,url):
        "获取文章内容"
        try:
            result = requests.get(url)
            if result is None: 
                raise Exception("获取Article失败[{0}]".format(url))
            if result.status_code != 200:
                raise Exception("获取Article失败[{0}][{1}]".format(url,result.status_code))
            "解析内容"
            soup = BeautifulSoup(result.content, 'html.parser')
            "标题 class=headline__text inline-placeholder vossi-headline-text"
            title = soup.select_one('h1[class="headline__text inline-placeholder vossi-headline-text"]')
            if title is None:
                raise Exception("获取文章标题失败[{0}]".format(url))
            title = title.text.strip()
            "内容 class=paragraph inline-placeholder vossi-paragraph"
            content = soup.select('p[class="paragraph inline-placeholder vossi-paragraph"]')
            if len(content) == 0:
                raise Exception("获取文章内容失败[{0}]".format(url))
            text_content = []
            for p in content:
                text_content.append(p.text.strip())
            #CSaveData("我是标题",result.content,"html")
            "3、保存文章"
            CSaveData(title,text_content,"docx")
            #
            print("获取文章完成[{0}]".format(url))
            
        except Exception as e:
            print(e)

if __name__ == '__main__':
    try:
        '''
        print(CSaveConfig.GetPathAndName())
        CSaveData("我是标题","我是内容:alfjdalkfjdlakjfalk!!!!!","docx")
        CSaveData("我是标题","我是内容:alfjdalkfjdlakjfalk!!!!!","txt")
        CSaveData("我是标题","我是内容:alfjdalkfjdlakjfalk!!!!!","html")
        '''
        cnn_url = "https://edition.cnn.com"
        CGetNews(cnn_url).Get()

    except Exception as e:
        print(e)